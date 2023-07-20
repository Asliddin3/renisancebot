from docx import Document
from docx.shared import Pt
import docx
import qrcode
from docx.shared import Cm


def update_doc_id(doc, id):
    for paragraph in doc.paragraphs:
        if "SHARTNOMA №________" == paragraph.text:
            paragraph.text = "SHARTNOMA №" + id
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(11)
            font.bold = True


def update_date(doc, year, day, month):
    string_to_replace = "Toshkent shahri 	2023 yil “_____” _____________"
    for paragraph in doc.paragraphs:
        if string_to_replace == paragraph.text:
            paragraph.text = "Toshkent shahri                                                    {}  yil “{}” {}".format(
                year, day, month)
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(10)
            font.bold = True


def update_contract(doc, contract_data):
    for row in doc.tables[0].rows:
        for key, value in contract_data.items():
            if row.cells[0].text == key:
                row.cells[1].text = value
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.runs[0]
                font = run.font
                font.size = Pt(10)
                font.bold = True


def update_name(doc, name):
    placeholder = "bir tomondan, ____________________________________________"
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph into parts
            parts = paragraph.text.split(placeholder)

            # Clear the paragraph
            paragraph.clear()

            # prepare the first run (text before the placeholder)
            before_run = paragraph.add_run(parts[0])
            # No need to set 'bold' for 'before_run' since it should remain normal

            # prepare the second run (substitute name)
            name_run = paragraph.add_run(name)
            name_run.font.size = Pt(10)  # Set the font size to 10
            name_run.bold = True  # make the name bold

            # If there is any text after the placeholder, add it as the third run
            if len(parts) > 1:
                after_run = paragraph.add_run(parts[1])
                name_run.font.size = Pt(10)  # Set the font size to 10

                # No need to set 'bold' for 'after_run' since it should remain normal


def update_price(doc, price, price_number):
    placeholder = "uchun to‘lov _______________ (______________________ ______________________________________________)"
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph into parts
            parts = paragraph.text.split(placeholder)

            # Clear the paragraph
            paragraph.clear()

            # prepare the first run (text before the placeholder)
            before_run = paragraph.add_run(parts[0])
            before_run.font.size = Pt(10)  # Set the font size to 10

            # prepare the second run (substitute price_number)
            price_number_run = paragraph.add_run(price_number)
            price_number_run.bold = True  # Make the price_number bold
            price_number_run.font.size = Pt(10)  # Set the font size to 10

            # prepare the third run (substitute price)
            price_run = paragraph.add_run(price)
            price_run.bold = True  # Make the price bold
            price_run.font.size = Pt(10)  # Set the font size to 10

            # If there is any text after the placeholder, add it as the fourth run
            if len(parts) > 1:
                after_run = paragraph.add_run(parts[1])
                after_run.font.size = Pt(10)  # Set the font size to 10


def update_student(doc, student_data):
    asd = 'Yashash manzili:______________________________\n____________________________________________'
    replacements = {
        "F.I.Sh.:______________________________________": student_data["name"],
        asd: student_data["address"],
        "Pasport ma’lumotlari:__________________________": student_data["passport"],
        "JSHSHIR:  ___________________________________": student_info["jshshir"],
        "Telefon raqami: ______________________________": student_data["number"],
        "Talaba _______   _____________________________": "Talaba ________ {}".format(student_data["name"])
    }
    table = doc.tables[1]
    cell = table.cell(0, 1)
    original_text = cell.text
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Talaba")
    run.bold = True
    for key, value in replacements.items():
        if key in original_text:
            original_text = original_text.replace(key, value)

    paragraph.add_run(original_text[7:])

    for run in paragraph.runs:
        font = run.font
        font.size = Pt(10)


contract_info = {
    "Ta’lim bosqichi:": "2-kurs",
    "Ta’lim shakli:": "Sirtqi",
    "O‘qish muddati:": "2-yil(2026)",
    "O‘quv kursi:": "Ingliz tili",
    "Ta’lim yo‘nalishi: ": "Moliya"
}
student_info = {
    "name": "F.I.Sh.: Asliddin Dehqonov",
    "address": "Yashash manzili: Toshkent viloyati Bekobod tumani Golang mahallasi 15-uy",
    "passport": "Pasport ma’lumotlari: FA1111111",
    "jshshir": "JSHSHIR:  21820719689744",
    "number": "Telefon raqami: 87 123 45 67",
}


def add_hyperlink(paragraph, url="https://t.me/renuadmisson", text="https://t.me/renuadmisson"):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)  # type: ignore

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')  # type: ignore
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )  # type: ignore

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')  # type: ignore

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')  # type: ignore

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = "\n\n" + text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def add_qr(doc, link="https://t.me/renuadmisson"):
    img = qrcode.make(link)
    img.save("qrcode.png")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('qrcode.png', width=Cm(3), height=Cm(3))


def update_student2(doc, student_data):
    replacements = {
        "F.I.Sh.:______________________________________": student_data["name"],
        "Yashash manzili:______________________________\n____________________________________________": student_data[
            "address"],
        "Pasport ma’lumotlari:__________________________": student_data["passport"],
        "JSHSHIR:  ___________________________________": student_data["jshshir"],
        "Telefon raqami: ______________________________": student_data["number"]
    }

    table = doc.tables[1]
    cell = table.cell(0, 1)
    old_text = cell.text
    cell.text = ''
    paragraph = cell.paragraphs[0]

    for key, replacement in replacements.items():
        # Split the paragraph into parts where the key is found
        parts = old_text.split(key)

        # Reconstruct the paragraph, inserting the student data formatted as bold
        if len(parts) >= 2:
            # The text before the key
            paragraph.add_run(parts[0])

            # Insert the student data
            student_data_run = paragraph.add_run(replacement)
            student_data_run.bold = True

            # The text after the key
            paragraph.add_run(parts[1])
            paragraph.add_run('\n')  # Break the paragraph after adding the student data
        else:
            # If the key was not found, add back the original text
            paragraph.add_run(old_text)
    # Fix the first line in the cell
    paragraph.runs[0].text = "Talaba:"
    paragraph.runs[0].bold = False
    paragraph.runs[0].font.size = Pt(10)  # Set the font size to 10

    for run in paragraph.runs[1:]:  # Skip the first run
        run.font.size = Pt(10)  # Set the font size to 10 for the other runs


# document = Document('shartnoma.docx')
#
# add_qr(document)
# add_hyperlink(document.add_paragraph())
# update_price(document, " (yigirma million besh yuz ellik ming)", "20550000")
# update_name(document, "Asliddin Dehqonov")
# update_doc_id(document, "5123")
# update_date(document, "2023", "20", "oktyabr")
# update_student(document, student_info)
# update_student2(document, student_info)
# update_contract(document, contract_info)
# document.save('output.docx')

def create_contract(data):
    document=Document("/root/univer-bot/renisancebot/generator/shartnoma.docx")
    # add_qr(document)
    add_hyperlink(document.add_paragraph())
    update_price(document, data["price_text"],data["price"])
    update_name(document, data["full_name"])
    update_doc_id(document, data["id"])
    update_date(document,data["year"], data["day"],data["month"])
    update_student(document, data["student_info"])
    # update_student2(document, data["student_info"])
    update_contract(document, data["contract_info"])
    add_qr(document,f'http://78.40.219.247:8000/contract/{data["id"]}/')
    # document.save(f'../documents/{data["id"]}/shartnoma.docx')
    document.save(f"/root/univer-bot/renisancebot/documents/{data['id']}/shartnoma.docx")





