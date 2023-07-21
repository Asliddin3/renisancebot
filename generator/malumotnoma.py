from PyPDF2.constants import UserAccessPermissions
from docx import Document
from docx.shared import Pt
import os
import qrcode
from docx.shared import Cm
import PyPDF2
import subprocess
info = {
    "name": "Asliddin Dehqonov ",
    "faculty": "moliya",
    "learn_type": "sirtqi",
    "id": "12421",
}

def update_payment_info(doc, infos):
    replacements = {
        "Abituriyent ______________________": f'Abituriyent {infos["name"]} ',
        "muassasasining __________": f'muassasasining {infos["faculty"]} ',
        "________ ta\'lim": f'{infos["learn_type"]} ta\'lim',
        "No ____________": f'No {infos["id"]}',
        "17.07.2023": infos["date"]
    }
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                run = paragraph.runs[0]
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(11)
                font.bold = False

    for shape in doc.inline_shapes:
        if shape.type == 3:  # Check if the shape is an image (type 3 corresponds to an image).
            shape = shape._element
            for key, value in replacements.items():
                rId = shape.getparent().xpath('./a:blip/@r:embed', namespaces=shape.nsmap)[0]
                img_part = doc.part.related_parts[rId]
                old_image_text = img_part._blob
                new_image_text = old_image_text.replace(key.encode(), value.encode())
                img_part._blob = new_image_text

# update_payment_info(document,info)
# document.save(f"documents/{info}/info.docx")
def add_qr(doc, link="https://t.me/renuadmisson"):
    img = qrcode.make(link)
    img.save("qrcode.png")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('qrcode.png', width=Cm(2), height=Cm(2))

def create_info(data):
    docx_file_path = "/root/univer-bot/renisancebot/generator/malumotnoma.docx"
    doc = Document(docx_file_path)
    update_payment_info(doc,data)
    add_qr(doc,link=f'http://78.40.219.247:8000/info/{data["path"]}/')
    add_qr(doc)
    directory_path = f"/root/univer-bot/renisancebot/documents/{data['path']}"
    os.makedirs(directory_path, exist_ok=True)
    # doc.save(f"../documents/{data['id']}/info.docx")
    doc.save(f"/root/univer-bot/renisancebot/documents/{data['path']}/malumotnoma.docx")
    subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
                    f"/root/univer-bot/renisancebot/documents/{data['path']}",
                    f"/root/univer-bot/renisancebot/documents/{data['path']}/malumotnoma.docx"])
    with open(f"/root/univer-bot/renisancebot/documents/{data['path']}/malumotnoma.pdf", 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)

        # Create a new PDF writer object
        pdf_writer = PyPDF2.PdfWriter()

        # Add each page from the original PDF to the writer object
        for page_num in range(len(pdf_reader.pages)):
            pdf_writer.add_page(pdf_reader.pages[page_num])

        # Set the permissions to disallow all modifications
        pdf_writer.add_blank_page()
        pdf_writer.encrypt('', 'umid1199', use_128bit=True, permissions_flag=0b0100)

        # Write the protected PDF to the output file
        with open(f"/root/univer-bot/renisancebot/documents/{data['path']}/malumotnoma.pdf", 'wb') as output_file:
            pdf_writer.write(output_file)
    # convert(f"/root/univer-bot/renisancebot/documents/{data['id']}/info.docx",
    #         f"/root/univer-bot/renisancebot/documents/{data['id']}/info.pdf")
    os.remove(f"/root/univer-bot/renisancebot/documents/{data['path']}/malumotnoma.docx")