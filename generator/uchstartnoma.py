import os

from docx import Document
from docx.shared import Pt
import docx
import qrcode
from docx.shared import Cm
from python_docx_replace import docx_replace
import subprocess

from docx2pdf import convert

def update_doc_id(doc, id):
    for paragraph in doc.paragraphs:
        if "SHARTNOMA №________" == paragraph.text:
            paragraph.text = "SHARTNOMA №" + id
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(11)
            font.bold = True


def update_date(doc, year, day, month):
    string_to_replace = "2023 yil “_____” _____________"
    for paragraph in doc.paragraphs:
        if string_to_replace in paragraph.text:
            paragraph.text = paragraph.text.replace(string_to_replace,
                                                    f"{year} yil “{day}” {month}".format(year, day, month))
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(10)
            font.bold = True


def update_contract(doc, contract_data):
    for row in doc.tables[0].rows:
        for key, value in contract_data.items():
            if row.cells[0].text == key:
                row.cells[1].text = value
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.runs[0]
                font = run.font
                font.size = Pt(10)
                font.bold = True


def update_student_name(doc, name):
    placeholder = "bir tomondan, ____________________ ______________________________________"
    splitPlace="____________________ ______________________________________"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph into parts
            parts = paragraph.text.split(splitPlace)

            # Clear the paragraph
            paragraph.clear()

            # prepare the first run (text before the placeholder)
            before_run = paragraph.add_run(parts[0])
            before_run.font.size = Pt(10)  # Set the font size to 10

            # No need to set 'bold' for 'before_run' since it should remain normal

            # prepare the second run (substitute name)
            name_run = paragraph.add_run(name + " ")
            name_run.font.size = Pt(10)  # Set the font size to 10
            # name_run.bold = True # make the name bold

            # If there is any text after the placeholder, add it as the third run
            if len(parts) > 1:
                after_run = paragraph.add_run(parts[1])
                name_run.font.size = Pt(10)  # Set the font size to 10

                # No need to set 'bold' for 'after_run' since it should remain normal


def update_sponsor_name(doc, name):
    placeholder = "ikkinchi tomondan, _________________________ ________________________________"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph into parts

            parts = paragraph.text.split(placeholder)

            # Clear the paragraph
            paragraph.clear()

            # prepare the first run (text before the placeholder)
            before_run = paragraph.add_run(parts[0])
            before_run.font.size = Pt(10)  # Set the font size to 10

            # No need to set 'bold' for 'before_run' since it should remain normal

            # prepare the second run (substitute name)
            name_run = paragraph.add_run(name)
            name_run.font.size = Pt(10)  # Set the font size to 10
            # name_run.bold = True # make the name bold

            # If there is any text after the placeholder, add it as the third run
            if len(parts) > 1:
                after_run = paragraph.add_run(parts[1])
                name_run.font.size = Pt(10)  # Set the font size to 10

                # No need to set 'bold' for 'after_run' since it should remain normal


def update_price(doc, price, price_number):
    placeholder = "to‘lov ________________________________________ (_____________________________________________________) "
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph into parts
            parts = paragraph.text.split(placeholder)

            # Clear the paragraph
            paragraph.clear()

            # prepare the first run (text before the placeholder)
            before_run = paragraph.add_run(parts[0])
            before_run.font.size = Pt(9)  # Set the font size to 9

            # prepare the second run (substitute price_number)
            price_number_run = paragraph.add_run(price_number)
            price_number_run.bold = True  # Make the price_number bold
            price_number_run.font.size = Pt(9)  # Set the font size to 9

            # prepare the third run (substitute price)
            price_run = paragraph.add_run(price)
            price_run.bold = True  # Make the price bold
            price_run.font.size = Pt(9)  # Set the font size to 9

            # If there is any text after the placeholder, add it as the fourth run
            if len(parts) > 1:
                after_run = paragraph.add_run(parts[1])
                after_run.font.size = Pt(10)  # Set the font size to 10


def update_student(doc, student_data):
    address = 'Yashash manzili:______________________ _____________________________________'
    replacements = {
        "F.I.Sh.: ______________________________ _____________________________________": student_data["name"],
        address: student_data["address"],
        "Pasport ma’lumotlari:_________________": student_data["passport"],
        "JSHSHIR:  __________________________": student_info["jshshir"],
        "Telefon raqami:_______________________": student_data["number"],
        "Talaba __________ ____________________": "Talaba(imzo) __________ {} ".format(student_data["name"])
    }
    table = doc.tables[1]
    cell = table.cell(0, 2)

    original_text = cell.text
    cell.text = ''

    for key, value in replacements.items():
        if key in original_text:
            original_text = original_text.replace(key, value)

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(original_text)

    font = run.font
    font.bold = False
    font.size = Pt(9)


def update_sponsor(doc, sponsor_data):
    replacements = {
        "Tashkilot nomi:_______________________ _____________________________________": sponsor_data["company_name"],
        "Manzil:______________________________ _____________________________________": sponsor_data["address"],
        "Bank:________________________________": sponsor_data["bank"],
        "X/r: _________________________________": sponsor_data["xr"],
        "Bank kodi: ___________________________": sponsor_data["bank_code"],
        "STIR: _______________________________": sponsor_data["stir"],
        "Telefon:______________________________": sponsor_data["number"],
        "Rahbar__________ ____________________": "Rahbar__________ {}".format(sponsor_data["ceo"]),
    }
    table = doc.tables[1]
    cell = table.cell(0, 1)

    original_text = cell.text
    cell.text = ''

    for key, value in replacements.items():
        if key in original_text:
            original_text = original_text.replace(key, value)

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(original_text)

    font = run.font
    font.bold = True
    font.size = Pt(9)


def add_hyperlink(paragraph, url="https://t.me/renuadmisson", text="https://t.me/renuadmisson"):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)  # type: ignore

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')  # type: ignore
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )  # type: ignore

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')  # type: ignore

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')  # type: ignore

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = "\n\n" + text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def add_qr(doc, link="https://t.me/renuadmisson"):
    img = qrcode.make(link)
    img.save("qrcode.png")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('qrcode.png', width=Cm(2), height=Cm(2))


contract_info = {
    "Ta’lim bosqichi:": "2-kurs",
    "Ta’lim shakli:": "Sirtqi",
    "O‘qish muddati:": "2-yil(2026)",
    "O‘quv kursi:": "Ingliz tili",
    "Ta’lim yo‘nalishi: ": "Moliya"
}

student_info = {
    "name": "F.I.Sh.: Asliddin Dehqonov",
    "address": "Yashash manzili: Toshkent viloyati Bekobod tumani Golang mahallasi 15-uy",
    "passport": "Pasport ma’lumotlari: FA1111111",
    "jshshir": "JSHSHIR:  21820719689744",
    "number": "Telefon raqami: 87 123 45 67",
}

sponsor_info = {
    "company_name": "Tashkilot nomi: Smart MChJ",
    "address": "Manzil: Toshkent viloyati Bekobod tumani Golang mahallasi 15-uy",
    "bank": "Bank: Mirobod bank",
    "xr": "X/r: 21820719689744",
    "bank_code": "Bank kodi: 14234",
    "stir": "STIR: stir123",
    "number": "Telefon: 71 123 45 67",
    "ceo": "Asqar Umirzoqov"
}

# document = Document('3shartnoma.docx')
# document.save('output.docx')
# update_doc_id(document, "5123")
# update_date(document, "2023", "20", "oktyabr")
# update_student_name(document, "Asliddin Dehqonov")
# update_sponsor_name(document, "Abror Dehqonov")
# update_contract(document, contract_info)
# update_price(document, " (yigirma million besh yuz ellik ming)", "20550000")
# update_student(document, student_info)
# update_sponsor(document, sponsor_info)
# add_hyperlink(document.add_paragraph())
# add_qr(document)
#
# document.save("output.docx")

def create_uchtamonlama(data):
    document = Document('/root/univer-bot/renisancebot/generator/uchshartnoma.docx')
    update_doc_id(document, data['id'])
    update_date(document, data["year"], data["day"],data["month"])
    update_student_name(document, data["full_name"])
    # update_sponsor_name(document, "Abror Dehqonov")
    update_contract(document, data["contract_info"])
    update_price(document, data["price_text"],data["price"])
    update_student(document, data["student_info"])
    # update_sponsor(document, sponsor_info)
    add_hyperlink(document.add_paragraph())
    add_qr(document,f'http://78.40.219.247:8000/document/{data["path"]}/')
    add_qr(document)
    # document.save(f"../documents/{data['id']}/uchtamonlama.docx")

    document.save(f"/root/univer-bot/renisancebot/documents/{data['path']}/uchtamonlama.docx")
    subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
                    f"/root/univer-bot/renisancebot/documents/{data['path']}",
                    f"/root/univer-bot/renisancebot/documents/{data['path']}/uchtamonlama.docx"])
    os.remove(f"/root/univer-bot/renisancebot/documents/{data['path']}/uchtamonlama.docx")

